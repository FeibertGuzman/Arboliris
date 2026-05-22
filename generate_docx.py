import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Sets background color of a cell (hex format, e.g. 'F1F5F9')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (in twentieths of a point, dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Sets fine borders for a cell (hex format, e.g., 'CBD5E1')."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, color in borders.items():
        if color:
            b_element = OxmlElement(f'w:{border_name}')
            b_element.set(qn('w:val'), 'single')
            b_element.set(qn('w:sz'), '4')  # fine line
            b_element.set(qn('w:space'), '0')
            b_element.set(qn('w:color'), color)
            tcBorders.append(b_element)
        else:
            b_element = OxmlElement(f'w:{border_name}')
            b_element.set(qn('w:val'), 'none')
            tcBorders.append(b_element)
            
    tcPr.append(tcBorders)

def format_run(run, font_name='Arial', size_pt=11, bold=False, italic=False, color_rgb=(15, 23, 42)):
    """Convenience helper to format font runs."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def add_header_styled_paragraph(doc, text, level):
    """Adds a heading with professional indigo/slate styling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    
    if level == 1:
        format_run(run, 'Arial', 16, bold=True, color_rgb=(79, 70, 229))  # Indigo
    elif level == 2:
        format_run(run, 'Arial', 13, bold=True, color_rgb=(30, 41, 59))   # Slate
    else:
        format_run(run, 'Arial', 11, bold=True, color_rgb=(79, 70, 229))
    return p

def create_doc():
    doc = docx.Document()
    
    # Document main metadata
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("TALENTO TECH")
    format_run(title_run, 'Arial', 24, bold=True, color_rgb=(79, 70, 229))
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Módulo: Aprendizaje Inteligente | Encuentro 1\nACTIVIDAD 1\n“De datos a decisiones: ML para transformar realidades”\nTaller colaborativo — Propuesta Iris ID3")
    format_run(sub_run, 'Arial', 12, italic=True, color_rgb=(100, 116, 139))
    
    # Table 0: Activity Instructions
    t0 = doc.add_table(rows=1, cols=1)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell0 = t0.rows[0].cells[0]
    set_cell_background(cell0, 'EEF2F6')  # Soft gray-blue background
    set_cell_margins(cell0, 200, 200, 250, 250)
    set_cell_borders(cell0, top='6366F1', bottom='6366F1', left='6366F1', right='6366F1')
    p0 = cell0.paragraphs[0]
    p0.add_run("Cómo funciona esta actividad:\n\n").bold = True
    p0.runs[0].font.color.rgb = RGBColor(79, 70, 229)
    run_inst = p0.add_run(
        "Su grupo diseñará una solución basada en Machine Learning orientada a resolver un problema real de clasificación.\n"
        "⚠️ Mentalidad del ejercicio: Eres arquitecto de soluciones de Inteligencia Artificial. Vas a estructurar:\n"
        "- El problema y variables a modelar.\n"
        "- Arquitectura de IA (Supervisado, Árbol ID3) e Ingesta de Datos (ETL).\n"
        "- Análisis Exploratorio (EDA) y DOFA de Datos.\n"
        "- Objetivos SMART y Dimensión Ética.\n\n"
        "Tiempo de diseño: 30 minutos | Presentación (elevator pitch): 5 minutos por grupo."
    )
    format_run(run_inst, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")  # Spacing
    
    # Table 1: Metadata Integrantes
    t1 = doc.add_table(rows=1, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    
    # Set widths for columns
    widths = [Inches(3.25), Inches(3.25)]
    for row in t1.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            set_cell_background(cell, 'F8FAFC')
            set_cell_margins(cell, 150, 150, 150, 150)
            set_cell_borders(cell, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
            
    c1_1 = t1.rows[0].cells[0]
    c1_1.paragraphs[0].add_run("Integrantes del grupo:\n").bold = True
    run_int1 = c1_1.paragraphs[0].add_run("• Feibert Alirio Guzmán Pérez")
    format_run(run_int1, 'Arial', 10, color_rgb=(71, 85, 105))
    
    c1_2 = t1.rows[0].cells[1]
    c1_2.paragraphs[0].add_run("Área disciplinar / Proyecto:\n").bold = True
    run_int2 = c1_2.paragraphs[0].add_run("Ingeniería de Sistemas / Ciencia de Datos e IA\nProyecto: Clasificación Especies de Iris (Algoritmo ID3)")
    format_run(run_int2, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Part 1: Design Template Header
    add_header_styled_paragraph(doc, "Parte 1 — Plantilla de diseño del proyecto (30 minutos)", 1)
    
    # Section 1: Problem Definition
    add_header_styled_paragraph(doc, "1. El problema que quieren resolver", 2)
    
    # Table 2: Section title
    t2 = doc.add_table(rows=1, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell2 = t2.rows[0].cells[0]
    set_cell_background(cell2, '6366F1')
    set_cell_margins(cell2, 100, 100, 150, 150)
    p2 = cell2.paragraphs[0]
    p2_run = p2.add_run("Descripción detallada del problema o necesidad identificada")
    format_run(p2_run, 'Arial', 11, bold=True, color_rgb=(255, 255, 255))
    
    # Table 3: Section content
    t3 = doc.add_table(rows=1, cols=1)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell3 = t3.rows[0].cells[0]
    set_cell_background(cell3, 'F8FAFC')
    set_cell_margins(cell3, 150, 150, 150, 150)
    set_cell_borders(cell3, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p3 = cell3.paragraphs[0]
    p3_run = p3.add_run(
        "La identificación manual de especies de flores de Iris (Setosa, Versicolor, Virginica) es un proceso ineficiente, lento y propenso a errores humanos para estudiantes de botánica y ecólogos aficionados. "
        "Dado que estas plantas comparten características morfológicas muy similares en sus sépalos y pétalos, la clasificación visual directa puede retrasar análisis biológicos. "
        "El proyecto propone automatizar esta labor mediante algoritmos de Machine Learning, convirtiendo dimensiones físicas de la flor en una etiqueta diagnóstica 100% confiable y explicable."
    )
    format_run(p3_run, 'Arial', 10.5, color_rgb=(30, 41, 59))
    
    doc.add_paragraph("")
    
    # Table 4: Variable objetivo
    doc.add_paragraph("🎯 Variable objetivo (target):").runs[0].bold = True
    t4 = doc.add_table(rows=1, cols=1)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell4 = t4.rows[0].cells[0]
    set_cell_background(cell4, 'F8FAFC')
    set_cell_margins(cell4, 150, 150, 150, 150)
    set_cell_borders(cell4, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p4 = cell4.paragraphs[0]
    p4_run = p4.add_run(
        "Variable objetivo: class (Especie de la flor).\n"
        "Tipo: categórica nominal (clasificación multiclase).\n"
        "Clases posibles: 'Iris-setosa', 'Iris-versicolor', 'Iris-virginica'.\n"
        "Cálculo: se determina a partir de las combinaciones de largo y ancho del sépalo y del pétalo del lirio."
    )
    format_run(p4_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 5: Research Question
    doc.add_paragraph("❓ Pregunta de investigación (orientada a ML):").runs[0].bold = True
    t5 = doc.add_table(rows=1, cols=1)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell5 = t5.rows[0].cells[0]
    set_cell_background(cell5, 'F8FAFC')
    set_cell_margins(cell5, 150, 150, 150, 150)
    set_cell_borders(cell5, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p5 = cell5.paragraphs[0]
    p5_run = p5.add_run(
        "¿Cómo construir y evaluar un clasificador supervisado de Árbol de Decisión ID3 (entropía) que prediga con una exactitud superior al 95% "
        "la especie de lirio Iris basándose en sus cuatro atributos morfológicos numéricos tabulares?"
    )
    format_run(p5_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 6: Impact and users
    doc.add_paragraph("🎯 ¿A quién impacta?").runs[0].bold = True
    t6 = doc.add_table(rows=1, cols=1)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell6 = t6.rows[0].cells[0]
    set_cell_background(cell6, 'F8FAFC')
    set_cell_margins(cell6, 150, 150, 150, 150)
    set_cell_borders(cell6, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p6 = cell6.paragraphs[0]
    p6_run = p6.add_run(
        "El proyecto beneficia directamente a estudiantes de ciencias naturales, docentes, clubes de biología y agrónomos. "
        "El impacto radica en agilizar la toma de decisión taxonómica en campo mediante un aplicativo web interpretable (Streamlit), "
        "explicando en lenguaje natural el proceso lógico del modelo."
    )
    format_run(p6_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Section 2: Architecture
    add_header_styled_paragraph(doc, "2. Arquitectura de IA y toma de decisión tecnológica", 2)
    
    # Table 8: Classification of ML approaches (supervised grid)
    t8 = doc.add_table(rows=1, cols=6)
    t8.alignment = WD_TABLE_ALIGNMENT.CENTER
    t8.autofit = False
    
    col_widths = [Inches(0.5), Inches(2.2), Inches(0.5), Inches(1.6), Inches(0.5), Inches(1.2)]
    for row in t8.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_background(cell, 'F8FAFC')
            set_cell_margins(cell, 100, 100, 100, 100)
            set_cell_borders(cell, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
            
    t8.rows[0].cells[0].paragraphs[0].add_run("[X]").bold = True
    format_run(t8.rows[0].cells[1].paragraphs[0].add_run(
        "MODELOS SUPERVISADOS (Clasificación)\nAplica porque contamos con un dataset balanceado con la especie (target) etiquetada. Modelo elegido: Decision Tree Classifier (entropía - ID3)."
    ), 'Arial', 9, color_rgb=(71, 85, 105))
    
    t8.rows[0].cells[2].paragraphs[0].add_run("[ ]").bold = True
    format_run(t8.rows[0].cells[3].paragraphs[0].add_run(
        "NO supervisados\nÚtiles para clustering descriptivo (agrupar flores por dimensiones morfológicas sin etiquetas)."
    ), 'Arial', 9, color_rgb=(71, 85, 105))
    
    t8.rows[0].cells[4].paragraphs[0].add_run("[ ]").bold = True
    format_run(t8.rows[0].cells[5].paragraphs[0].add_run(
        "RL / Preentrenados\nNo aplica a este tipo de problema estructurado."
    ), 'Arial', 9, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 9: Selected Model info
    doc.add_paragraph("🎯 Modelo seleccionado:").runs[0].bold = True
    t9 = doc.add_table(rows=1, cols=1)
    t9.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell9 = t9.rows[0].cells[0]
    set_cell_background(cell9, 'F8FAFC')
    set_cell_margins(cell9, 150, 150, 150, 150)
    set_cell_borders(cell9, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p9 = cell9.paragraphs[0]
    p9_run = p9.add_run(
        "Modelo seleccionado: DecisionTreeClassifier (algoritmo ID3).\n"
        "Criterio de partición: Entropía ('entropy') y ganancia de información.\n"
        "Parámetros clave: random_state=1, división estratificada 80/20 train/test.\n"
        "Entrada: sepallength, sepalwidth, petallength, petalwidth. Salida: especie."
    )
    format_run(p9_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 10: Justification
    doc.add_paragraph("🧠 Justificación:").runs[0].bold = True
    t10 = doc.add_table(rows=1, cols=1)
    t10.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell10 = t10.rows[0].cells[0]
    set_cell_background(cell10, 'F8FAFC')
    set_cell_margins(cell10, 150, 150, 150, 150)
    set_cell_borders(cell10, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p10 = cell10.paragraphs[0]
    p10_run = p10.add_run(
        "Los árboles de decisión operan dividiendo iterativamente el set de datos para purificar los subnodos. "
        "En este proyecto, se determinó que el ancho del pétalo (petal width) ofrece el mayor grado de ganancia de información inicial, "
        "pudiendo clasificar directamente la clase Iris-setosa sin condiciones adicionales. El aplicativo Streamlit aprovecha esta rápida inferencia l-ogica."
    )
    format_run(p10_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Insert Tree visual representation in model section
    if os.path.exists("image/Arbol.png"):
        doc.add_paragraph("Visualización de la Estructura de Decisión (ID3):").runs[0].bold = True
        doc.add_picture("image/Arbol.png", width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph("")
    
    # Table 11: Why is it the best option
    doc.add_paragraph("❓ ¿Por qué es la mejor opción?").runs[0].bold = True
    t11 = doc.add_table(rows=1, cols=1)
    t11.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell11 = t11.rows[0].cells[0]
    set_cell_background(cell11, 'F8FAFC')
    set_cell_margins(cell11, 150, 150, 150, 150)
    set_cell_borders(cell11, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p11 = cell11.paragraphs[0]
    p11_run = p11.add_run(
        "Un árbol de decisión ID3 es un modelo de 'caja blanca' (explicable). A diferencia de redes neuronales, el usuario puede inspeccionar "
        "las reglas lógicas escritas (ej: si petalwidth <= 0.8 cm es Setosa). Esto genera transparencia absoluta y facilita la asimilación del "
        "modelo por parte de botánicos no programadores."
    )
    format_run(p11_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Section 3: Data and ETL
    add_header_styled_paragraph(doc, "Datos y analítica para la toma de decisiones", 2)
    add_header_styled_paragraph(doc, "🔄 Proceso ETL (Extract, Transform, Load)", 3)
    
    # Table 12: Data needed and sources
    t12 = doc.add_table(rows=3, cols=2)
    t12.alignment = WD_TABLE_ALIGNMENT.CENTER
    t12.autofit = False
    
    t12.rows[0].cells[0].merge(t12.rows[0].cells[1])
    t12.rows[0].cells[0].paragraphs[0].text = "3. Los datos que necesitarían"
    t12.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(t12.rows[0].cells[0], 'EEF2F6')
    set_cell_margins(t12.rows[0].cells[0], 100, 100, 150, 150)
    
    set_cell_background(t12.rows[1].cells[0], 'F1F5F9')
    t12.rows[1].cells[0].paragraphs[0].text = "¿Qué tipo de datos necesitarían?"
    t12.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    
    set_cell_background(t12.rows[1].cells[1], 'F1F5F9')
    t12.rows[1].cells[1].paragraphs[0].text = "¿De dónde los obtendrían?"
    t12.rows[1].cells[1].paragraphs[0].runs[0].bold = True
    
    for r in [1, 2]:
        for c in [0, 1]:
            set_cell_margins(t12.rows[r].cells[c], 120, 120, 120, 120)
            set_cell_borders(t12.rows[r].cells[c], top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
            
    t12.rows[2].cells[0].width = Inches(3.25)
    t12.rows[2].cells[1].width = Inches(3.25)
    
    run_dat1 = t12.rows[2].cells[0].paragraphs[0].add_run(
        "Se requieren las mediciones físicas en cm de las flores:\n"
        "• sepallength (4.3 a 7.9 cm)\n"
        "• sepalwidth (2.0 a 4.4 cm)\n"
        "• petallength (1.0 a 6.9 cm)\n"
        "• petalwidth (0.1 a 2.5 cm)\n"
        "El dataset contiene 150 instancias balanceadas de lirios Iris."
    )
    format_run(run_dat1, 'Arial', 9.5, color_rgb=(71, 85, 105))
    
    run_dat2 = t12.rows[2].cells[1].paragraphs[0].add_run(
        "Se extrae del repositorio oficial de OpenML (ID: 61, dataset_61_iris.arff).\n"
        "La fuente es histórica y confiable (Ronald Fisher, 1936), ideal para calibrar y validar algoritmos de clasificación multiclase."
    )
    format_run(run_dat2, 'Arial', 9.5, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 13: Pipeline ML summary
    doc.add_paragraph("Pipeline ML de Ingesta y Modelado:").runs[0].bold = True
    t13 = doc.add_table(rows=1, cols=1)
    t13.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell13 = t13.rows[0].cells[0]
    set_cell_background(cell13, 'F8FAFC')
    set_cell_margins(cell13, 150, 150, 150, 150)
    set_cell_borders(cell13, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p13 = cell13.paragraphs[0]
    p13_run = p13.add_run(
        "Pipeline de Ingesta:\n"
        "1. Conectar origen OpenML -> 2. Ingesta por Pandas (Extract) -> 3. Limpieza de duplicados/nulos (Transform) -> "
        "4. Partición train_test_split estratificada (80/20) -> 5. Entrenamiento DecisionTreeClassifier(entropy) -> "
        "6. Inferencia y despliegue en Streamlit (Load)."
    )
    format_run(p13_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Insert Infographic in ETL section
    if os.path.exists("image/Infografia.png"):
        doc.add_paragraph("Flujo del Ciclo de Vida del Machine Learning:").runs[0].bold = True
        doc.add_picture("image/Infografia.png", width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph("")
    
    # Table 14: Extract
    doc.add_paragraph("🔹 Extract — ¿De dónde obtienen los datos?").runs[0].bold = True
    t14 = doc.add_table(rows=1, cols=1)
    t14.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell14 = t14.rows[0].cells[0]
    set_cell_background(cell14, 'F8FAFC')
    set_cell_margins(cell14, 150, 150, 150, 150)
    set_cell_borders(cell14, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p14 = cell14.paragraphs[0]
    p14_run = p14.add_run(
        "Se extraen dinámicamente mediante el comando pd.read_csv(url) desde la fuente de OpenML. "
        "Esta carga automatizada asegura que el dataset esté fresco, sea reproducible y cuente con la cabecera correcta."
    )
    format_run(p14_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 15: Transform
    doc.add_paragraph("🔹 Transform — ¿Cómo preparan los datos?").runs[0].bold = True
    t15 = doc.add_table(rows=1, cols=1)
    t15.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell15 = t15.rows[0].cells[0]
    set_cell_background(cell15, 'F8FAFC')
    set_cell_margins(cell15, 150, 150, 150, 150)
    set_cell_borders(cell15, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p15 = cell15.paragraphs[0]
    p15_run = p15.add_run(
        "El dataset Iris es limpio por naturaleza, sin embargo, se realiza:\n"
        "• Validación física de rangos continuos.\n"
        "• Separación de la matriz X (medidas) y vector y (especie).\n"
        "• Separación estratificada mediante train_test_split para conservar la distribución balanceada de las tres clases en test y train."
    )
    format_run(p15_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 16: Load
    doc.add_paragraph("🔹 Load — ¿Cómo dejan los datos listos?").runs[0].bold = True
    t16 = doc.add_table(rows=1, cols=1)
    t16.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell16 = t16.rows[0].cells[0]
    set_cell_background(cell16, 'F8FAFC')
    set_cell_margins(cell16, 150, 150, 150, 150)
    set_cell_borders(cell16, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p16 = cell16.paragraphs[0]
    p16_run = p16.add_run(
        "Carga del modelo entrenado y de las variables estructuradas directamente en la memoria caché del aplicativo Streamlit. "
        "Esto permite que la aplicación invoque al clasificador de forma instantánea al detectar cambios en los sliders del usuario."
    )
    format_run(p16_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Section 4: EDA
    add_header_styled_paragraph(doc, "🔍 ANÁLISIS EXPLORATORIO (EDA)", 2)
    
    t18 = doc.add_table(rows=1, cols=1)
    t18.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell18 = t18.rows[0].cells[0]
    set_cell_background(cell18, 'F8FAFC')
    set_cell_margins(cell18, 150, 150, 150, 150)
    set_cell_borders(cell18, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p18 = cell18.paragraphs[0]
    p18_run = p18.add_run(
        "Antes del entrenamiento, se analiza descriptivamente el dataset:\n"
        "• Conteo de clases (50 Setosa, 50 Versicolor, 50 Virginica).\n"
        "• Búsqueda de valores atípicos y nulos mediante resúmenes dt.describe() y df.isnull().sum().\n"
        "• Visualización de la dispersión de las flores respecto al largo y ancho de pétalos, donde se aprecia una separación espacial lineal para Setosa."
    )
    format_run(p18_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Table 20: DOFA Matrix
    doc.add_paragraph("Matriz DOFA — Gestión de Datos en Proyectos de IA:").runs[0].bold = True
    t20 = doc.add_table(rows=5, cols=5)
    t20.alignment = WD_TABLE_ALIGNMENT.CENTER
    t20.autofit = False
    
    dofa_widths = [Inches(1.0), Inches(1.4), Inches(1.4), Inches(1.4), Inches(1.4)]
    
    dofa_headers = ['Factor', 'Fortalezas (+)', 'Oportunidades (+)', 'Debilidades (-)', 'Amenazas (-)']
    for i, header_text in enumerate(dofa_headers):
        cell = t20.rows[0].cells[i]
        set_cell_background(cell, 'E2E8F0')
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_run = p.add_run(header_text)
        format_run(p_run, 'Arial', 9.5, bold=True, color_rgb=(30, 41, 59))
        
    dofa_data = [
        ['Calidad de datos', 'Dataset balanceado, limpio y sin nulos.', 'Cargar imágenes de lirios reales.', 'Bajo volumen de muestras (150).', 'Mutación botánica o microclimas.'],
        ['Acceso', 'Disponible públicamente en OpenML.', 'Crear app web accesible en móvil.', 'Requiere conexión para cargar datos.', 'API de OpenML caída temporalmente.'],
        ['Integración', 'Fácil de consumir con scikit-learn.', 'Integrar con herramientas de campo.', 'El modelo solo lee 4 dimensiones.', 'Formatos de datos alterados.'],
        ['Sesgo', 'Balance absoluto en las clases.', 'Ampliar con especies híbridas.', 'Mayor error en bordes difusos.', 'Variabilidad morfológica inusual.']
    ]
    
    for row_idx, row_text_list in enumerate(dofa_data):
        row = t20.rows[row_idx + 1]
        for col_idx, text in enumerate(row_text_list):
            cell = row.cells[col_idx]
            set_cell_background(cell, 'F8FAFC')
            set_cell_margins(cell, 100, 100, 100, 100)
            set_cell_borders(cell, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
            p = cell.paragraphs[0]
            p_run = p.add_run(text)
            if col_idx == 0:
                format_run(p_run, 'Arial', 9, bold=True, color_rgb=(71, 85, 105))
            else:
                format_run(p_run, 'Arial', 9, color_rgb=(100, 116, 139))
                
    for row in t20.rows:
        for i, cell in enumerate(row.cells):
            cell.width = dofa_widths[i]
            
    doc.add_paragraph("")
    
    # Section 5: SMART Goals
    add_header_styled_paragraph(doc, "🎯 OBJETIVOS SMART", 2)
    
    t22 = doc.add_table(rows=1, cols=1)
    t22.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell22 = t22.rows[0].cells[0]
    set_cell_background(cell22, 'F8FAFC')
    set_cell_margins(cell22, 150, 150, 150, 150)
    set_cell_borders(cell22, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p22 = cell22.paragraphs[0]
    p22_run = p22.add_run(
        "Objetivo General SMART:\n"
        "Diseñar, entrenar y desplegar durante el módulo del taller un clasificador de Árbol de Decisión ID3 (entropía) sobre el dataset Iris, "
        "obteniendo una exactitud global mayor al 95% y empaquetándolo en una interfaz interactiva de Streamlit comprensible para botánicos."
    )
    format_run(p22_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Section 6: Specific Objectives
    doc.add_paragraph("✅ Objetivos Específicos:").runs[0].bold = True
    
    specific_goals = [
        "1. Problema: Estructurar la clasificación morfológica manual ineficiente, reduciendo errores botánicos.",
        "2. Modelo ML: Entrenar un DecisionTreeClassifier regulando hiperparámetros para evitar el sobreajuste.",
        "3. Datos: Limpiar y preparar las 150 instancias de lirios de OpenML usando un pipeline ETL documentado.",
        "4. Riesgos: Controlar el traslape en la frontera de versicolor y virginica mediante validación cruzada y evaluación de matriz de confusión.",
        "5. Presentación: Levantar y desplegar un aplicativo Streamlit intuitivo para la visualización y estimación de especie."
    ]
    
    for goal in specific_goals:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(goal)
        format_run(run, 'Arial', 10, color_rgb=(71, 85, 105))
        
    doc.add_paragraph("")
    
    # Section 7: Ethics
    add_header_styled_paragraph(doc, "5. La dimensión ética del proyecto", 2)
    
    t29 = doc.add_table(rows=6, cols=4)
    t29.alignment = WD_TABLE_ALIGNMENT.CENTER
    t29.autofit = False
    
    eth_widths = [Inches(0.6), Inches(2.2), Inches(3.2), Inches(0.5)]
    
    eth_headers = ['Marca', 'Consideración ética', 'Cómo lo abordarían?', 'Aplica?']
    for i, text in enumerate(eth_headers):
        cell = t29.rows[0].cells[i]
        set_cell_background(cell, 'E2E8F0')
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        p_run = p.add_run(text)
        format_run(p_run, 'Arial', 9.5, bold=True, color_rgb=(30, 41, 59))
        if i == 0 or i == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    eth_data = [
        ['[X]', 'Privacidad y protección de datos', 'No se colectan datos personales, solo medidas físicas florales.', 'Sí'],
        ['[X]', 'Sesgo en los datos', 'Vigilar el balance del set (50 de cada clase) para que no favorezca una especie.', 'Sí'],
        ['[X]', 'Transparencia y explicabilidad', 'Se expone visualmente el árbol y sus reglas lógicas en la app y el HTML.', 'Sí'],
        ['[X]', 'Riesgo de etiquetar', 'Se evitan clasificaciones sesgadas, definiendo límites claros.', 'Sí'],
        ['[X]', 'Supervisión humana', 'El usuario (botánico/estudiante) toma la decisión final basada en el modelo.', 'Sí']
    ]
    
    for r_idx, row_list in enumerate(eth_data):
        row = t29.rows[r_idx + 1]
        for c_idx, text in enumerate(row_list):
            cell = row.cells[c_idx]
            set_cell_background(cell, 'F8FAFC')
            set_cell_margins(cell, 100, 100, 100, 100)
            set_cell_borders(cell, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
            p = cell.paragraphs[0]
            p_run = p.add_run(text)
            if c_idx == 0 or c_idx == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                format_run(p_run, 'Arial', 9, bold=True, color_rgb=(79, 70, 229))
            else:
                format_run(p_run, 'Arial', 9, color_rgb=(71, 85, 105))
                
    for row in t29.rows:
        for i, cell in enumerate(row.cells):
            cell.width = eth_widths[i]
            
    doc.add_paragraph("")
    
    # Section 8: Elevator Pitch
    add_header_styled_paragraph(doc, "Parte 2 — Guia para el elevator pitch (5 minutos por grupo)", 1)
    
    t30 = doc.add_table(rows=5, cols=2)
    t30.alignment = WD_TABLE_ALIGNMENT.CENTER
    t30.autofit = False
    
    pitch_widths = [Inches(1.2), Inches(5.3)]
    for r_idx in range(5):
        row = t30.rows[r_idx]
        for c_idx in range(2):
            cell = row.cells[c_idx]
            set_cell_background(cell, 'F8FAFC')
            set_cell_margins(cell, 100, 100, 100, 100)
            set_cell_borders(cell, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
            
    pitch_structure = [
        ['Min 0-1', 'Problema: La identificación botánica manual es ineficiente y propensa a errores por similitud física de flores Iris.'],
        ['Min 1-2', 'Solución: Un clasificador de Árbol de Decisión ID3 que infiere la especie con explicabilidad y rapidez.'],
        ['Min 2-3', 'Datos: 150 registros de OpenML cargados dinámicamente mediante pipeline ETL balanceado.'],
        ['Min 3-4', 'Ética: Garantías de anonimato, explicabilidad visual del árbol de decisión y toma de decisión en manos del botánico.'],
        ['Min 4-5', 'Impacto: Exactitud global de 96.67%, integración interactiva en Streamlit para agilizar la labor educativa y científica en campo.']
    ]
    
    for r_idx, val_list in enumerate(pitch_structure):
        row = t30.rows[r_idx]
        p_c0 = row.cells[0].paragraphs[0]
        p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_run(p_c0.add_run(val_list[0]), 'Arial', 9.5, bold=True, color_rgb=(79, 70, 229))
        
        p_c1 = row.cells[1].paragraphs[0]
        format_run(p_c1.add_run(val_list[1]), 'Arial', 9.5, color_rgb=(71, 85, 105))
        
    for row in t30.rows:
        for i, cell in enumerate(row.cells):
            cell.width = pitch_widths[i]
            
    doc.add_paragraph("")
    
    # Section 9: Evaluation Rubric
    add_header_styled_paragraph(doc, "Rubrica de evaluación entre pares", 1)
    
    t31 = doc.add_table(rows=6, cols=5)
    t31.alignment = WD_TABLE_ALIGNMENT.CENTER
    t31.autofit = False
    
    rub_widths = [Inches(1.8), Inches(1.4), Inches(1.4), Inches(1.4), Inches(0.5)]
    
    rub_headers = ['Criterio', 'Excelente (3 pts)', 'Bueno (2 pts)', 'En desarrollo (1 pt)', 'Ptos']
    for i, text in enumerate(rub_headers):
        cell = t31.rows[0].cells[i]
        set_cell_background(cell, 'E2E8F0')
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        p_run = p.add_run(text)
        format_run(p_run, 'Arial', 9, bold=True, color_rgb=(30, 41, 59))
        
    rub_data = [
        ['1. Pertinencia del problema', 'Problema real, taxonómico, de clasificación botánica claro.', 'Relevante, pero algo general.', 'No definido o irrelevante.', '/ 3'],
        ['2. Justificación técnica', 'Arquitectura ID3 (entropía) justificada con base en ganancia de información.', 'Arquitectura correcta pero justificación vaga.', 'Incorrecta o sin justificación.', '/ 3'],
        ['3. Consideración ética', 'Identifica riesgos de sesgo y explicabilidad, aportando mitigaciones.', 'Identifica un riesgo con mitigación genérica.', 'Sin reflexión ética.', '/ 2'],
        ['4. Enlace a GitHub / App', 'Carga exitosa del modelo y la app de Streamlit.', 'Carga parcial.', 'Carga fallida o semi-parcial.', '/ 2'],
        ['TOTAL:', 'TOTAL:', 'TOTAL:', 'TOTAL:', '/ 10']
    ]
    
    for r_idx, row_list in enumerate(rub_data):
        row = t31.rows[r_idx + 1]
        for c_idx, text in enumerate(row_list):
            cell = row.cells[c_idx]
            set_cell_background(cell, 'F8FAFC')
            set_cell_margins(cell, 80, 80, 80, 80)
            set_cell_borders(cell, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
            p = cell.paragraphs[0]
            p_run = p.add_run(text)
            if r_idx == 4 or c_idx == 0:
                format_run(p_run, 'Arial', 8.5, bold=True, color_rgb=(30, 41, 59))
            else:
                format_run(p_run, 'Arial', 8.5, color_rgb=(71, 85, 105))
                
    for row in t31.rows:
        for i, cell in enumerate(row.cells):
            cell.width = rub_widths[i]
            
    doc.add_paragraph("")
    
    # Section 10: Greatest Challenge and Mitigation
    add_header_styled_paragraph(doc, "✍️ Análisis del equipo", 2)
    
    doc.add_paragraph("¿Cuál es el mayor desafío en su proyecto?").runs[0].bold = True
    t32 = doc.add_table(rows=1, cols=1)
    t32.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell32 = t32.rows[0].cells[0]
    set_cell_background(cell32, 'F8FAFC')
    set_cell_margins(cell32, 120, 120, 150, 150)
    set_cell_borders(cell32, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p32 = cell32.paragraphs[0]
    p32_run = p32.add_run(
        "El mayor desafío radica en el solapamiento morfológico natural entre las especies Iris-versicolor e Iris-virginica en el rango "
        "intermedio de dimensiones del pétalo. En esta frontera, las flores presentan dimensiones casi idénticas, lo que dificulta la "
        "separación lineal del árbol y genera una pequeña tasa de clasificaciones incorrectas (Entropía residual)."
    )
    format_run(p32_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    doc.add_paragraph("¿Cómo lo mitigarían estratégicamente?").runs[0].bold = True
    t33 = doc.add_table(rows=1, cols=1)
    t33.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell33 = t33.rows[0].cells[0]
    set_cell_background(cell33, 'F8FAFC')
    set_cell_margins(cell33, 120, 120, 150, 150)
    set_cell_borders(cell33, top='E2E8F0', bottom='E2E8F0', left='E2E8F0', right='E2E8F0')
    p33 = cell33.paragraphs[0]
    p33_run = p33.add_run(
        "La mitigación se enfoca en tres frentes: 1) Ajustar la profundidad del árbol (max_depth) para evitar el sobreajuste en el ruido. "
        "2) Presentar probabilidades de clasificación en lugar de solo etiquetas duras en Streamlit, transparentando la incertidumbre en la frontera. "
        "3) Considerar en etapas futuras la recolección de variables físicas secundarias adicionales (textura de la hoja, color exacto del pétalo) "
        "para facilitar la diferenciación."
    )
    format_run(p33_run, 'Arial', 10, color_rgb=(71, 85, 105))
    
    doc.add_paragraph("")
    
    # Section 11: Quantitative validation images
    doc.add_paragraph("Validación del Desempeño del Modelo (Métricas y Confusión):").runs[0].bold = True
    
    if os.path.exists("image/Matriz Confusión.png"):
        doc.add_picture("image/Matriz Confusión.png", width=Inches(3.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if os.path.exists("image/Comparación Valoras Predichos.png"):
        doc.add_picture("image/Comparación Valoras Predichos.png", width=Inches(3.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph("")
    
    # Bibliografía
    add_header_styled_paragraph(doc, "Bibliografía", 1)
    
    bib_items = [
        "1. Fisher, R. A. The use of multiple measurements in taxonomic problems. Annals of Eugenics, 1936.",
        "2. Quinlan, J. R. Induction of decision trees. Induction of decision trees. Machine Learning, 1986.",
        "3. Scikit-learn Developers. DecisionTreeClassifier and classification metrics documentation. Scikit-learn official docs.",
        "4. Guzmán Pérez, F. A. Árbol de Decisión ID3 con el Dataset Iris. Repositorio del proyecto IA Talento Tech, 2026."
    ]
    
    for bib in bib_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(bib)
        format_run(run, 'Arial', 10, color_rgb=(71, 85, 105))
        
    # Save the document overwriting the template
    save_path = r"c:\Users\Feibert\OneDrive\Documents\AGENTE\Arboles\Proyecto\Material_1_Proyecto_IA_Talento_Tech.docx"
    doc.save(save_path)
    print("Success: Generated docx file at:", save_path)

if __name__ == "__main__":
    create_doc()
